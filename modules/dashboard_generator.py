from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
from reportlab.lib import colors
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.charts.linecharts import HorizontalLineChart
from reportlab.graphics.charts.piecharts import Pie
import matplotlib.pyplot as plt
import io
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import tempfile

class DashboardGenerator:
    """Generates dashboard reports in PDF and Word formats"""
    
    def create_report(self, df, format='pdf'):
        """Create dashboard report in specified format"""
        if format.lower() == 'pdf':
            return self._create_pdf_report(df)
        else:
            return self._create_word_report(df)
    
    def _create_pdf_report(self, df):
        """Create PDF dashboard report"""
        # filename = '/home/claude/dashboard_report.pdf'
        temp_dir = tempfile.gettempdir()
        filename = os.path.join(temp_dir, 'dashboard_report.pdf')
        
        doc = SimpleDocTemplate(filename, pagesize=letter)
        story = []
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#0066CC'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#0066CC'),
            spaceAfter=12
        )
        
        # Title
        story.append(Paragraph('Analytics Dashboard Report', title_style))
        story.append(Paragraph(f'Generated on {datetime.now().strftime("%B %d, %Y")}', styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Executive Summary
        story.append(Paragraph('Executive Summary', heading_style))
        
        summary_data = [
            ['Metric', 'Value'],
            ['Total Records', f'{len(df):,}'],
            ['Total Columns', str(len(df.columns))],
            ['Data Completeness', f'{(1 - df.isnull().sum().sum() / df.size) * 100:.1f}%'],
            ['Duplicate Records', str(df.duplicated().sum())]
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0066CC')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(summary_table)
        story.append(Spacer(1, 0.5*inch))
        
        # Data Quality Metrics
        story.append(Paragraph('Data Quality Analysis', heading_style))
        
        missing_by_col = df.isnull().sum()
        missing_cols = missing_by_col[missing_by_col > 0]
        
        if len(missing_cols) > 0:
            story.append(Paragraph(f'<b>Missing Values Found:</b> {len(missing_cols)} columns have missing data', styles['Normal']))
            
            missing_data = [['Column', 'Missing Count', 'Percentage']]
            
            for col in missing_cols.head(5).index:
                count = missing_cols[col]
                pct = (count / len(df)) * 100
                missing_data.append([col, str(count), f'{pct:.1f}%'])
            
            missing_table = Table(missing_data, colWidths=[2.5*inch, 1.5*inch, 1.5*inch])
            missing_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FF6B6B')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(Spacer(1, 0.2*inch))
            story.append(missing_table)
        else:
            story.append(Paragraph('✓ No missing values detected', styles['Normal']))
        
        story.append(Spacer(1, 0.3*inch))
        
        # Numeric Statistics
        story.append(Paragraph('Statistical Summary', heading_style))
        
        numeric_cols = df.select_dtypes(include=['number']).columns
        
        if len(numeric_cols) > 0:
            stats_data = [['Column', 'Mean', 'Median', 'Std Dev', 'Min', 'Max']]
            
            for col in numeric_cols[:5]:
                stats_data.append([
                    col,
                    f'{df[col].mean():.2f}',
                    f'{df[col].median():.2f}',
                    f'{df[col].std():.2f}',
                    f'{df[col].min():.2f}',
                    f'{df[col].max():.2f}'
                ])
            
            stats_table = Table(stats_data, colWidths=[1.5*inch, 1*inch, 1*inch, 1*inch, 1*inch, 1*inch])
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0066CC')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(stats_table)
        
        story.append(Spacer(1, 0.5*inch))
        
        # Generate charts
        story.append(Paragraph('Visual Analytics', heading_style))
        
        # Bar chart
        if len(numeric_cols) > 0:
            chart_img = self._create_bar_chart(df, numeric_cols[0])
            if chart_img:
                story.append(Image(chart_img, width=5*inch, height=3*inch))
                story.append(Spacer(1, 0.3*inch))
        
        # Recommendations
        story.append(PageBreak())
        story.append(Paragraph('Recommendations', heading_style))
        
        recommendations = self._generate_recommendations(df)
        
        for rec in recommendations:
            story.append(Paragraph(f'• {rec}', styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story)
        
        return filename
    
    def _create_word_report(self, df):
        """Create Word dashboard report"""
        # filename = '/home/claude/dashboard_report.docx'
        temp_dir = tempfile.gettempdir()
        filename = os.path.join(temp_dir, 'dashboard_report.docx')
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Analytics Dashboard Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph(f'Generated on {datetime.now().strftime("%B %d, %Y")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        summary_data = [
            ('Total Records', f'{len(df):,}'),
            ('Total Columns', str(len(df.columns))),
            ('Data Completeness', f'{(1 - df.isnull().sum().sum() / df.size) * 100:.1f}%'),
            ('Duplicate Records', str(df.duplicated().sum()))
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.rows[i].cells[0].text = label
            summary_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # Data Quality
        doc.add_heading('Data Quality Analysis', 1)
        
        missing_by_col = df.isnull().sum()
        missing_cols = missing_by_col[missing_by_col > 0]
        
        if len(missing_cols) > 0:
            doc.add_paragraph(f'Missing Values Found: {len(missing_cols)} columns have missing data')
            
            missing_table = doc.add_table(rows=len(missing_cols[:5]) + 1, cols=3)
            missing_table.style = 'Light Grid Accent 1'
            
            # Headers
            missing_table.rows[0].cells[0].text = 'Column'
            missing_table.rows[0].cells[1].text = 'Missing Count'
            missing_table.rows[0].cells[2].text = 'Percentage'
            
            for i, col in enumerate(missing_cols.head(5).index, 1):
                count = missing_cols[col]
                pct = (count / len(df)) * 100
                
                missing_table.rows[i].cells[0].text = col
                missing_table.rows[i].cells[1].text = str(count)
                missing_table.rows[i].cells[2].text = f'{pct:.1f}%'
        else:
            doc.add_paragraph('✓ No missing values detected')
        
        doc.add_paragraph()
        
        # Numeric Statistics
        doc.add_heading('Statistical Summary', 1)
        
        numeric_cols = df.select_dtypes(include=['number']).columns
        
        if len(numeric_cols) > 0:
            stats_table = doc.add_table(rows=len(numeric_cols[:5]) + 1, cols=6)
            stats_table.style = 'Light Grid Accent 1'
            
            # Headers
            headers = ['Column', 'Mean', 'Median', 'Std Dev', 'Min', 'Max']
            for i, header in enumerate(headers):
                stats_table.rows[0].cells[i].text = header
            
            # Data
            for i, col in enumerate(numeric_cols[:5], 1):
                stats_table.rows[i].cells[0].text = col
                stats_table.rows[i].cells[1].text = f'{df[col].mean():.2f}'
                stats_table.rows[i].cells[2].text = f'{df[col].median():.2f}'
                stats_table.rows[i].cells[3].text = f'{df[col].std():.2f}'
                stats_table.rows[i].cells[4].text = f'{df[col].min():.2f}'
                stats_table.rows[i].cells[5].text = f'{df[col].max():.2f}'
        
        doc.add_page_break()
        
        # Recommendations
        doc.add_heading('Recommendations', 1)
        
        recommendations = self._generate_recommendations(df)
        
        for rec in recommendations:
            doc.add_paragraph(rec, style='List Bullet')
        
        # Save document
        doc.save(filename)
        
        return filename
    
    def _create_bar_chart(self, df, column):
        """Create bar chart using matplotlib"""
        try:
            fig, ax = plt.subplots(figsize=(6, 4))
            
            # Get top 10 values
            if df[column].nunique() < 20:
                data = df[column].value_counts().head(10)
            else:
                data = df.groupby(df.index)[column].sum().head(10)
            
            data.plot(kind='bar', ax=ax, color='#0066CC')
            ax.set_title(f'{column} Distribution')
            ax.set_xlabel('Category')
            ax.set_ylabel('Value')
            
            plt.tight_layout()
            
            # Save to buffer
            buf = io.BytesIO()
            plt.savefig(buf, format='png', dpi=100)
            buf.seek(0)
            plt.close()
            
            return buf
            
        except Exception as e:
            print(f"Error creating chart: {str(e)}")
            return None
    
    def _generate_recommendations(self, df):
        """Generate data quality recommendations"""
        recommendations = []
        
        # Missing values
        missing_count = df.isnull().sum().sum()
        if missing_count > 0:
            recommendations.append(f'Address {missing_count} missing values to improve data completeness')
        
        # Duplicates
        dup_count = df.duplicated().sum()
        if dup_count > 0:
            recommendations.append(f'Remove {dup_count} duplicate records to ensure data integrity')
        
        # Data types
        object_cols = df.select_dtypes(include=['object']).columns
        if len(object_cols) > 0:
            recommendations.append(f'Review {len(object_cols)} text columns for potential standardization')
        
        # Outliers
        numeric_cols = df.select_dtypes(include=['number']).columns
        
        for col in numeric_cols:
            try:
                Q1 = df[col].quantile(0.25)
                Q3 = df[col].quantile(0.75)
                IQR = Q3 - Q1
                outliers = df[(df[col] < (Q1 - 1.5 * IQR)) | (df[col] > (Q3 + 1.5 * IQR))]
                
                if len(outliers) > 0:
                    recommendations.append(f'Investigate {len(outliers)} potential outliers in {col}')
                    break  # Only report once
            except:
                pass
        
        if not recommendations:
            recommendations.append('Data quality looks good! No immediate issues detected.')
        
        return recommendations
